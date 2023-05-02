from email import message
from Company.models import Job_Profiles, Companies,skills,industry
from django.contrib import messages 

from django.shortcuts import render, redirect, reverse
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import DetailView
from django.contrib.auth.models import User
from django.shortcuts import render
from django.http import HttpResponse,JsonResponse
from Company.models import Companies,Job_Profiles,shortList
from resume.models import Candidate,Course,Skill,Work_Experience,Projects
from exams.models import ExamResult,Questions,Options,Answers,ExamDuration
from docx import Document
import os
from pathlib import Path
from django.conf import settings


my_file =os.path.join(settings.BASE_DIR,"Company\ProfilesInfo\.Job_Profiles object (1).docx")
print(my_file)
# my_file = os.path.join(base_dir,".Job_Profiles object (1).docx")



from docx.shared import Inches
#..............ML.............#
import pandas as pd
# For performing text preprocessing
from nltk.stem import PorterStemmer, LancasterStemmer, SnowballStemmer, WordNetLemmatizer
from nltk.corpus import stopwords
from string import punctuation
import re 
import joblib
import pickle
import docx2txt
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# f=open(my_file,"r")
# print(f.read())   
# Create your views here.

def company_home(request):
    if not Candidate.objects.filter(username=request.user.username):
        return render(request,'company_dashboard.html')
    return render(request,'resume/candiHome.html')


#company registration
def company_registration(request):
    if request.method== 'POST':
        print(request)
        company_name=request.POST.get('cname')
        company_mobile_no=request.POST.get('mobileno')
        company_email=request.POST.get('email')
        companyuserid=request.POST.get('userid')
        state=request.POST.get('state')
        city=request.POST.get('city')
        password1=request.POST.get('pass1')
        password2=request.POST.get('pass1')


        if Companies.objects.filter(companyuserid=companyuserid):
            messages.error(request,"Company already exit please try diffrent id")
            print("1")
            return redirect('company_registration')
        if Companies.objects.filter(company_email=company_email):
            messages.error(request,"Company email id is already exit!")
            print("2")
            return redirect('company_registration')
        if password1 != password2:
            messages.error(request, "Passwords didnt match!")
            print("3")
            return redirect('company_registration')
            
        mycompany=Companies(company_name=company_name, company_contact_no=company_mobile_no,company_email=company_email,companyuserid=companyuserid,state=state,city=city,pass1=password1,pass2=password2)
        mycompany.save()
        u=User(username=companyuserid,email=company_email,password=password1)
        u.set_password(password1)
        u.save()

        return redirect('login')
    return render(request,'companyregistrationPage.html')

# def company_login(request):
#     if request.method=='POST':
#         print(request)
#         username=request.get('username')
#         password=request.get('pass1')
#         if(Companies.objects.filter(username=username)):
#             comp=Companies.objects.get(username=username)
#             pas=comp.pass1
#             if(password==pas):
#                 return render(request,'add.html')

def companyhome(request):
    print(request)
    user=request.user.username
    if not Candidate.objects.filter(username=user):
        comp=Companies.objects.get(companyuserid=user)
        # info=Job_Profiles.objects.values('id','profile_name','job_info','no_of_vacancies')
        info=Job_Profiles.objects.filter(company_id=comp)
        indus=industry.objects.all()
        applicants=0
        for i in info:
            temp=ExamResult.objects.filter(jobId=i.id)
            applicants+=len(temp)
        count=len(info)
        return render(request,'company_dashboard.html',{'company_info':info,'count':count,'applicants':applicants,'indus':indus})
    return redirect('candihome')
    # return render(request,'resume/candiHome.html')



def ADD(request):

    if request.method=='POST':
        
        job_role=request.POST.get('jobrole')
        industry=request.POST.get('industry')
        job_des=request.POST.get('jobdes')
        vacancies=request.POST.get('vacancies')
        # job_skills=request.POST.get('skills')
        salary=request.POST.get('salary')
        experience=request.POST.get('experience')
        req_qual=request.POST.get('req_qual')
        location=request.POST.get('location')
        
        keys=list((request.POST).keys())
        print(keys)
        company_id=Companies.objects.get(companyuserid=request.user.username)
        #job_profile_id=Job_Profiles.objects.get(id=i_d)
        
        print(company_id)
        document=Document()
        document.add_heading('Job Role:',1)
        document.add_paragraph(job_role)
        document.add_heading('Industry', level=1)
        document.add_paragraph(industry)
        document.add_heading('Job Description:', level=1)
        document.add_paragraph(job_des)
        document.add_heading('No of Vacancies:',level=1)
        document.add_paragraph(vacancies)
        document.add_heading('salary:',level=1)
        document.add_paragraph(salary)
        document.add_heading('Experience:',level=1)
        document.add_paragraph(experience)
        document.add_heading('Required Qualification',level=1)
        document.add_paragraph(req_qual)

        document.add_heading('Required Skills:', level=1)

        data=Job_Profiles(profile_name=job_role,company_id=company_id,job_info=job_des,salary=salary,no_of_vacancies=vacancies,education=req_qual,industry=industry,joblocation_address=location,experience=experience)
        data.save()
        id=data.id
        job_profile_id=Job_Profiles.objects.get(id=id)
            
        for key in keys:
            temp=str(key)
            skill='skill'
            if skill in temp:
                skill=request.POST.getlist(temp)[0]
                document.add_paragraph(skill, style='List Number')
                data2=skills(skills=skill,company_id=company_id,job_profile_id=job_profile_id)
                data2.save()
        
        # print(PROJECT_ROOT)

        filepath=r"C:\Users\ACER\Downloads\Recruitment_management2\Recruitment_management2\Company\ProfilesInfo\."+str(job_profile_id)+".docx"
        # path="Company\ProfilesInfo\."+str(id)+".docx"
        # document.save(os.path.join(settings.BASE_DIR,path))
        document.save(filepath)
        return redirect('comphome')
    else:
        info=Job_Profiles.objects.values('pk','profile_name','job_info','no_of_vacancies')
        print("maiiii")
        print(info)
        return render(request,'company_dashboard.html',{'company_info':info})

def delete(request,id):
    role=Job_Profiles.objects.get(id=id)
    print(role)
    role.delete()
    return redirect('comphome')


def job_info(request):
    return render(request,'views.html')

class JobDetailView(LoginRequiredMixin, DetailView):
    model = Job_Profiles
    def get_context_data(self,**kwargs):
        context = super(JobDetailView, self).get_context_data(**kwargs)
        print(self.kwargs['pk'])
        id=Job_Profiles.objects.get(id=self.kwargs['pk'])
        print("KKKKKK")
        Test=Questions.objects.filter(jobId=id)
        TotalScore=0
        for t in Test:
            TotalScore+=t.score
        print(TotalScore)
        sorts=[]
        results=ExamResult.objects.filter(jobId=id).order_by('-totalScore')
        # rfilepath=r""
        job_description=docx2txt.process(my_file)
        path="Company\ProfilesInfo\."+str(id)+".docx"
        job_description = docx2txt.process(os.path.join(settings.BASE_DIR,path))
        for i in results:
            path1="resume\ResumeFiles\."+str(i.candidateId)+".docx"
            resume = docx2txt.process(os.path.join(settings.BASE_DIR,path1))
            # resume = docx2txt.process(rfilepath+'\resume\ResumeFiles\\'+ str(i.candidateId)+'.docx')
            resume=docx2txt.process(my_file)

            content = [job_description, resume]
            cv = CountVectorizer()
            count_matrix = cv.fit_transform(content)
            mat = cosine_similarity(count_matrix)
            sorts.append(tuple([i,str(mat[1][0]*100),i.totalScore]))
        sorts.sort(key = lambda x: (x[2],x[1]),reverse=True)
        print(sorts)
        # print(skills.objects.filter(job_profile_id=id))
        context['skills']=skills.objects.filter(job_profile_id=id)
        # context['testres']=ExamResult.objects.filter(jobId=id).order_by('-totalScore')
        context['testres']=sorts
        context['num']=len(ExamResult.objects.filter(jobId=id))
        context['TotalScore']=TotalScore
        context['jobId']=id.id
        return context

def disresume(request):
    print(request.POST)
    if request.method=='POST':
        candi=request.POST['id']
        candidate=Candidate.objects.get(id=candi)
        skills=list(Skill.objects.filter(candidateId=Candidate.objects.get(id=candi)).values())
        # return HttpResponse({'bool':True})
        workExp=list(Work_Experience.objects.filter(candidateId=Candidate.objects.get(id=candi)).values())
        courses=list(Course.objects.filter(candidateId=Candidate.objects.get(id=candi)).values())
        prjs=list(Projects.objects.filter(candidateId=Candidate.objects.get(id=candi)).values())
        return JsonResponse({'username':candidate.candidate_name,'college':candidate.college,'cgpa':candidate.cgpa,'skills':skills,'workExp':workExp,'courses':courses,'prjs':prjs})



def accept(request):
    print(request.POST)
    if request.method=='POST':
        print(request.POST['id'])
        id=request.POST['id']
        jobId=request.POST['jobId']
        check=request.POST['check']
        print(check)
        print(id)
        print(jobId)
        c=Candidate.objects.get(id=id)
        j=Job_Profiles.objects.get(id=jobId)
        if check=="1":
            s=shortList(candidate_id=c,job_id=j)
            s.save()
            s=ExamResult.objects.get(candidateId=c,jobId=j)
            s.status=True
            s.save()
            return JsonResponse({"bool":True})
        else:
            s=shortList.objects.get(candidate_id=c,job_id=j)
            s.delete()
            s=ExamResult.objects.get(candidateId=c,jobId=j)
            s.status=False
            s.save()
            return JsonResponse({"bool":False})


def deltest(request,id):
    Questions.objects.filter(jobId=id).delete()
    ExamDuration.objects.filter(jobId=id).delete()
    # questions.delete()
    return redirect('job_details', id)
