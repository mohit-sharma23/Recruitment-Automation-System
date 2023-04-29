from django.shortcuts import render,redirect
from resume.models import resumedata,Candidate,Work_Experience,Course,Skill,Projects,Resume
from Company.models import Job_Profiles,skills,shortList
from exams.models import ExamResult
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from docx import Document
from docx.shared import Inches
import datetime
#..............ML.............#
import numpy as np
import pandas as pd
# For performing text preprocessing
from nltk.stem import PorterStemmer, LancasterStemmer, SnowballStemmer, WordNetLemmatizer
from nltk.corpus import stopwords
from string import punctuation
import re
from sklearn.feature_extraction.text import TfidfVectorizer
# For encoding the targets
from sklearn.preprocessing import LabelEncoder
# For creating train, test and validation sets
from sklearn.model_selection import train_test_split
# For implementing machine learning models

from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier,GradientBoostingClassifier, AdaBoostClassifier
from sklearn.svm import SVC 
from xgboost import XGBClassifier
import xgboost as xgb
# For evaluating the models
from sklearn.metrics import classification_report, accuracy_score
import joblib
import pickle
from sklearn.impute import SimpleImputer


# Create your views here.

def candi_regis(request):
    if request.method =="POST":
        print(request.POST)
        candidate_name=request.POST['name']
        candidate_email=request.POST['email']
        pass1=request.POST['pass1']
        pass2=request.POST['pass2']
        username=request.POST['username']
        phno=request.POST['phone']
        address=request.POST['address']
        college=request.POST['college']
        cgpa=request.POST['cgpa']


        if Candidate.objects.filter(username=username):
            messages.error(request,"Candidate already exit please try diffrent id")
            print("1")
            return redirect('candi_regis')
        if Candidate.objects.filter(candidate_email=candidate_email):
            messages.error(request,"Candidate email id is already exit!")
            print("2")
            return redirect('candi_regis')
        if pass1 != pass2:
            messages.error(request, "Passwords didnt match!")
            print("3")
            return redirect('candi_regis')

        candidate=Candidate(candidate_name=candidate_name,candidate_email=candidate_email,pass1=pass1,pass2=pass2,username=username,phno=phno,address=address,college=college,cgpa=cgpa)
        candidate.save()
        u=User(username=username,email=candidate_email,password=pass1)
        u.set_password(pass1)
        u.save()
        return redirect('login')

    return render(request,'resume/candiRegist.html')

def candihome(request):
    user=request.user.username
    candidate=Candidate.objects.get(username=user)
    CandiSkills=Skill.objects.filter(candidateId=candidate)
    FiltSkills=[]
    for s in CandiSkills:
        JobSkills=skills.objects.filter(skills=s.skill)
        for i in JobSkills:
            FiltSkills.append(i.job_profile_id)
    # print(FiltSkills)
    FSkills=set(FiltSkills)
    print(FSkills)
    job=Job_Profiles.objects.all()
    resume=Resume.objects.filter(candidateId=candidate)
    # print(resume)
    num_can=len(Candidate.objects.all())
    num_res=len(Resume.objects.all())
    num_placed=len(ExamResult.objects.filter(status=True))
    jobs={
        'job':job,
        'resume':resume,
        'num_can':num_can,
        'num_res':num_res,
        'FSkills':FSkills,
        'num_placed':num_placed
    }
    # if(skills):
        # print(skills)
    return render(request,'resume/candiHome.html',jobs)

def create_resume(request):
    
    if request.method == "POST":
        print(request.POST)
        keys=list((request.POST).keys())
        print(keys)
        user=Candidate.objects.get(username=request.user.username)
        for key in keys:
            temp=str(key)
            companyName=''
            WorkDetails=''
            Internship=False
            comp='comp'
            course='course'
            prj='prj'
            skill='skill'
            if comp in temp:
                companyName=request.POST.getlist(temp)[0]
                WorkDetails=request.POST.getlist(temp)[1]
                print(companyName)
                print(WorkDetails)
                if request.POST.getlist(temp)[2]=='Yes':
                    Internship=True
                else:
                    Internship=False
                workExp=Work_Experience(candidateId=user,companyName=companyName,workDetails=WorkDetails,internship=Internship)
                workExp.save()
            if course in temp:
                course=request.POST.getlist(temp)[0]
                platform=request.POST.getlist(temp)[1]
                courses=Course(candidateId=user,course=course,platform=platform)
                courses.save()
            if prj in temp:
                print(request.POST.getlist(temp))
                prj_name=request.POST.getlist(temp)[0]
                prj_link=request.POST.getlist(temp)[1]
                prj_des=request.POST.getlist(temp)[2]
                project=Projects(candidateId=user,prj_name=prj_name,prj_link=prj_link,prj_des=prj_des)
                project.save()
            if skill in temp:
                skill=request.POST.getlist(temp)[0]
                s=Skill(candidateId=user,skill=skill)
                s.save()
        CanRes=Resume.objects.filter(candidateId=user)
        print(CanRes)
        if CanRes.count() == 0:
            resume=Resume(candidateId=user)
            resume.save()
        document=Document()
        document.add_heading(user.candidate_name,level=0)
        document.add_paragraph(user.candidate_email)
        document.add_paragraph(user.phno)
        document.add_paragraph(user.address)
        document.add_paragraph(user.college)
        document.add_paragraph('CGPA: '+user.cgpa)
        document.add_heading('Courses:',level=1)
        allcourses=Course.objects.filter(candidateId=user)
        for c in allcourses:
            document.add_heading(c.course,level=3)
            document.add_paragraph(c.platform)
        document.add_heading('Projects:',level=1)
        allprjs=Projects.objects.filter(candidateId=user)
        for p in allprjs:
            document.add_heading(p.prj_name,level=3)
            document.add_paragraph(p.prj_link)
            document.add_paragraph(p.prj_des)
        document.add_heading('Skills:',level=1)
        allskills=Skill.objects.filter(candidateId=user)
        for s in allskills:
            document.add_paragraph(s.skill)
        document.add_heading('Work Experience:',level=1)
        allexps=Work_Experience.objects.filter(candidateId=user)
        for ex in allexps:
            document.add_heading(ex.companyName,level=3)
            document.add_paragraph(ex.workDetails)
            if ex.internship:
                document.add_paragraph('Type- Internship')
            else:
                document.add_paragraph('Type- Job')
        filepath=r"C:\Users\ACER\Downloads\Recruitment_management2\Recruitment_management2\resume\ResumeFiles\."+str(user)+".docx"
        document.save(filepath)
        return redirect('candihome')

    # else:
    allSkills=skills.objects.values('skills').distinct()
    # print(allSkills)
    context={
        'skills':allSkills,
    }
    return render(request, 'resume/createresume.html',context)

def candi_profile(request):
    candi=Candidate.objects.get(username=request.user.username)
    return render(request,'resume/candi_profile.html',{'candi':candi})

@login_required(login_url='login')
def job_info(request,id):
    job=Job_Profiles.objects.get(id=id)
    comp=job.company_id.company_name
    user=request.user.username
    candidate=Candidate.objects.get(username=user)
    resume=Resume.objects.filter(candidateId=candidate)
    s=skills.objects.filter(job_profile_id=job)
    print(resume)
    context={
        'job':job,
        'comp':comp,
        'resume':resume,
        's':s,
    }
    return render(request,'resume/jobdetail.html',context)

def update_resume(request):
    user=Candidate.objects.get(username=request.user.username)
    wrkExp=Work_Experience.objects.filter(candidateId=user)
    courses=Course.objects.filter(candidateId=user)
    prjs=Projects.objects.filter(candidateId=user)
    s=Skill.objects.filter(candidateId=user)
    allSkills=skills.objects.values('skills').distinct()

    context={
        'wrkExp':wrkExp,
        'courses':courses,
        'prjs':prjs,
        'skills':s,
        'allskills':allSkills,

    }
    return render(request,'resume/updateresume.html',context)

def del_exp(request):
    user=Candidate.objects.get(username=request.user.username)
    print('hello')
    Work_Experience.objects.filter(candidateId=user).delete()
    return redirect('update_resume')

def del_course(request):
    user=Candidate.objects.get(username=request.user.username)
    print('hello')
    Course.objects.filter(candidateId=user).delete()
    return redirect('update_resume')

def del_prj(request):
    user=Candidate.objects.get(username=request.user.username)
    print('hello')
    Projects.objects.filter(candidateId=user).delete()
    return redirect('update_resume')

def del_skill(request):
    user=Candidate.objects.get(username=request.user.username)
    print('hello')
    Skill.objects.filter(candidateId=user).delete()
    return redirect('update_resume')


def inter_offer(request,id):
    job=Job_Profiles.objects.get(id=id)
    context={
        'job':job,
    }
    return render(request,'resume/interview_offer.html',context)


# RECOMMENDATIONS

    
def apply_job(request):
    user=request.user.username
    candidate=Candidate.objects.get(username=user)
    jobs=ExamResult.objects.filter(candidateId=candidate)
    preprocess = Preprocess()
    data = pd.read_csv('static/naukri_com-job_sample.csv')
    # data=data.iloc[0:10000]
    data.head()
    data.isnull().sum()[data.isnull().sum()>0]
    data.drop(['numberofpositions','site_name'],axis=1,inplace=True)
    # to_fill = ['education', 'skills']
    # for col in to_fill:
    #     imputer = SimpleImputer(strategy='most_frequent')
    #     # print(data[col])
    #     # data[[col]]=imputer.fit(data[[col]])
    #     data[col]=imputer.fit_transform(data[[col]])
    data.drop(['jobid','uniq_id'],axis=1,inplace=True)
    data = data[data['experience'] != 'Not Mentioned']
    experience_lower = []
    experience_upper = []
    invalid = []
    for idx, row in data.iterrows():
        try:
            text = re.sub('yrs', '', row['experience'])
            splits = text.split('-')
            experience_lower.append(int(splits[0]))
            experience_upper.append(int(splits[1]))
        except:
            invalid.append(row['experience'])
    data = data[~data['experience'].isin(invalid)]
    data['experience_lower'] = data['experience'].apply(lambda x: int(x.split('-')[0]))
    data['experience_upper'] = data['experience'].apply(lambda x: int(re.sub('yrs','', x.split('-')[1])))

    data.drop(['experience'], axis=1, inplace=True)
    data['postdate'] = data['postdate'].astype(str).apply(lambda x: x[:-5])
    data['job_age']=datetime.datetime.today() - pd.to_datetime(data['postdate'])
    data['job_age'] = data['job_age'].dt.days
    replacements = {
        'joblocation_address': {
            r'(Bengaluru/Bangalore)': 'Bangalore',
            r'Bengaluru': 'Bangalore',
            r'Hyderabad / Secunderabad': 'Hyderabad',
            r'Mumbai , Mumbai': 'Mumbai',
            r'Noida': 'NCR',
            r'Delhi': 'NCR',
            r'Gurgaon': 'NCR', 
            r'Delhi/NCR(National Capital Region)': 'NCR',
            r'Delhi , Delhi': 'NCR',
            r'Noida , Noida/Greater Noida': 'NCR',
            r'Ghaziabad': 'NCR',
            r'Delhi/NCR(National Capital Region) , Gurgaon': 'NCR',
            r'NCR , NCR': 'NCR',
            r'NCR/NCR(National Capital Region)': 'NCR',
            r'NCR , NCR/Greater NCR': 'NCR',
            r'NCR/NCR(National Capital Region) , NCR': 'NCR', 
            r'NCR , NCR/NCR(National Capital Region)': 'NCR', 
            r'Bangalore , Bangalore / Bangalore': 'Bangalore',
            r'Bangalore , karnataka': 'Bangalore',
            r'NCR/NCR(National Capital Region)': 'NCR',
            r'NCR/Greater NCR': 'NCR',
            r'NCR/NCR(National Capital Region) , NCR': 'NCR'
        }
    }
    data.replace(replacements, regex=True, inplace=True)
    data['industry'] = data['industry'].astype(str).apply(lambda x: x.split('/')[0])
    data['industry'] = data['industry'].str.strip()
    # source: https://www.kaggle.com/code/danala26/data-cleaning-naukri-com
    # data['Education'] = data['education'].str.split(' ')
    # data['Education'] = data['Education'].apply(lambda x: x[1] if len(x) > 1 else x[0])

    # data['Education'] = data['Education'].replace(('B.Tech/B.E.','Graduation','Other','-','Not','B.Tech/B.E.,','Postgraduate',
    #                                                'PG:CA','Diploma,','B.Com,','B.Pharma,','B.A,','BCA,','B.Sc,','MBA/PGDM','B.B.A,',
    #                                               'PG:Other','Doctorate:Doctorate','Post'),
    #                                               ('B.Tech','B.Tech','B.Tech','B.Tech','B.Tech','B.Tech','B.Tech',
    #                                               'CA','Diploma','B.Com','B.Pharma','B.A','BCA','B.Sc','MBA','BBA',
    #                                               'B.Tech','Doctorate','B.Tech'))

    # data['Skills'] = data['skills'].str.split(" - ")
    # data['Skills'] = data['Skills'].apply(lambda x: x[1] if len(x) > 1 else x[0])
    # Removing jobs which have less than 10 postings as they are very rare and affect model performance
    majority_industries = data['industry'].value_counts()[data['industry'].value_counts()>=10].index
    data = data[data['industry'].isin(majority_industries)]
    data.isnull().sum()[data.isnull().sum()>0]

    # joblocation_imputer = SimpleImputer(strategy='most_frequent')
    # data['joblocation_address'] = joblocation_imputer.fit_transform(data[['joblocation_address']])
    # jobage_imputer = SimpleImputer(strategy='mean')
    # data['job_age'] = jobage_imputer.fit_transform(data[['job_age']])

    preprocess.fit_transform(data['jobdescription'],data['industry'])
    # random_job = data.sample(n=1,random_state=1234)
    # # random_job=data2.iloc[0:1]
    # # data
    # random_job
    model2 = xgb.XGBClassifier()
    model2.load_model('static/new_model.json')
    # model2
    data = pd.read_csv('static/naukri_com-job_sample.csv')
    # data=data.iloc[0:10000]
    random_job = data.sample(n=1,random_state=1234)
    # random_job=data2.iloc[0:1]
    # data
    # random_job
    z=preprocess.transform(random_job)
    # pred=model2.predict(z)
    d=[24,24,24,24,24,24,24,24,24,24,24,24,24,24]
    decoded_pred = preprocess.encoder.inverse_transform(d)[0]

    # decoded_pred = preprocess.encoder.inverse_transform(pred)[0]
    # print(decoded_pred)
    # decoded_pred='IT-Software'
    recomm=Job_Profiles.objects.filter(industry=decoded_pred)
    
    context={
        'jobs':jobs,
        'candidate':candidate,
        'recomm':recomm
    }
    return render(request,'resume/jobs.html',context)

class Preprocess:
    def __init__(self,method='WordNetLemmatizer'):
        # WordNetLemmatizer is recommended because it reduces the given word to the root word 
        # by referring to the WordNet corpus unlike other stemming techniques which just 
        # truncate the word by removing the suffix, which is why I have set it as default
        self.method = method
        self.stemmers = {
            'PorterStemmer':PorterStemmer(),
            'LancasterStemmer':LancasterStemmer(),
            'SnowballStemmer':SnowballStemmer(language='english'),
            'WordNetLemmatizer':WordNetLemmatizer()
        }
        self.stemmer = self.stemmers[self.method]
        # Remove punctuation signs and stopwords for better results 
        self.stopWords = list(punctuation) + list(stopwords.words('english'))
        # Adding custom stopwords for better preprocessing, feel free to add more
        self.moreStopWords = ['job','description','requirement','skill', 'qualification']
        self.stopWords.extend(self.moreStopWords)
        self.encoder = LabelEncoder()
        # Using tf-idf vectorizer because it not only relies on the count but also the 
        # number of documents it occurs in 
        # tf * log(N/df), where tf = term frequency/count of words 
        # N = total number of documents 
        # df = document frequency (number of documents containing that word)
        # Count vectorizer gets tricked by the term frequency but in tf-idf it does not happen
        # eg - if the word occurs frequently in almost all documents,it may be a filler word 
        # which was ignored in stopwords, so it can trick the count vectorizer,but in tf-idf
        # N/df almost = 1, so log(N/df) will be almost 0 and hence the word will not be given 
        # much importance which is desirable, hence we should use tfidf vectorizer instead of 
        # count vectorizer 
        self.vectorizer = TfidfVectorizer()
        self.isFitted = False
    def preprocess(self, message):
        message = message.lower()
        #Remove links 
        message = re.sub('http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+#]|[!*\(\),]|'\
                        '(?:%[0-9a-fA-F][0-9a-fA-F]))+','', message)
        # Remove extra spaces 
        message = re.sub(' +', ' ', message)
        # Remove mentions 
        message =re.sub("(@[A-Za-z0-9_]+)","", message)
        # Remove Hashtags
        message = re.sub('#[A-Za-z0-9_]+','', message)
        # Remove all non alphanumeric characters 
        message = re.sub("^[A-Za-z0-9_-]*$", "", message)
        # Remove Emojis 
        emoji_pattern = re.compile(
            "["
            u"\U0001F600-\U0001F64F"  # emoticons
            u"\U0001F300-\U0001F5FF"  # symbols & pictographs
            u"\U0001F680-\U0001F6FF"  # transport & map symbols
            u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
            u"\U00002702-\U000027B0"
            u"\U000024C2-\U0001F251"
            "]+",
            flags=re.UNICODE,
        )
        message = emoji_pattern.sub('',message)
        if self.method == 'WordNetLemmatizer':
            message = ' '.join([self.stemmer.lemmatize(word) for word in message.split() if word not in self.moreStopWords])
        else:
            message = ' '.join([self.stemmer.stem(word) for word in message.split() if word not in self.moreStopWords])
        return message 
    def fit(self,X,y=None):
        self.vectorizer.fit(X)
        if y is not None:
            self.encoder.fit(y)
        self.isFitted=True
    def transform(self, X, y=None):
        self.fit(X)
        if not self.isFitted:
            raise NotImplementedError('Please fit first by calling the fit function')
        X = self.vectorizer.transform(X)
        if y is not None:
            y = self.encoder.transform(y)
            return X,y
        else:
            return X 
    def fit_transform(self,X,y=None):
        self.fit(X,y)
        X,y = self.transform(X,y)
        return X,y
    